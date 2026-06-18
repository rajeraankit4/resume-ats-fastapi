import os
import re
from typing import List, Optional, Tuple

import fitz  # PyMuPDF
import numpy as np
import psycopg2
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import docx
except ImportError:
    print("python-docx not found: pip install python-docx")
    docx = None


# ── Environment ──────────────────────────────────────────────────────────────

load_dotenv()
DATABASE_URL = os.getenv("DATABASE_URL")

# ── Model (singleton) ────────────────────────────────────────────────────────

_model = None
_model_load_attempted = False


def get_model() -> Optional[SentenceTransformer]:
    global _model, _model_load_attempted
    if _model is not None:
        return _model
    if _model_load_attempted:
        return None
    _model_load_attempted = True
    try:
        _model = SentenceTransformer("all-MiniLM-L6-v2", local_files_only=True)
        return _model
    except Exception as e:
        print(f"SentenceTransformer unavailable, falling back to TF-IDF: {e}")
        return None


# ── Domain definitions ───────────────────────────────────────────────────────

DOMAINS: dict[str, str] = {
    "Software": (
        "Software engineering, backend, frontend, full stack, QA, test automation, "
        "DevOps, platform engineering, cloud infrastructure, site reliability, "
        "cybersecurity, application development, APIs, web development, databases."
    ),
    "Analyst": (
        "Data analyst, business analyst, research analyst, data science, machine learning, "
        "artificial intelligence, reporting, dashboards, business intelligence, forecasting, "
        "statistics, data visualization, model building, experimentation."
    ),
    "Core_NonTech": (
        "Core engineering, embedded systems, firmware, hardware, electronics, electrical, "
        "mechanical, manufacturing, automotive, operations, procurement, supply chain, "
        "sales, marketing, HR, finance, administration, non-technical roles."
    ),
}

DOMAIN_SKILL_TERMS = {
    "Software": {
        "software", "developer", "backend", "frontend", "full stack", "web", "api", "rest api",
        "graphql", "node.js", "react", "angular", "vue", "express", "javascript", "typescript",
        "html", "css", "java", "python", "c++", "sql", "mongodb", "postgresql", "redis",
        "git", "github", "docker", "aws", "gcp", "azure", "linux", "kubernetes", "terraform",
        "jenkins", "github actions", "devops", "cloud", "microservices", "system design",
        "qa", "selenium", "cypress", "jest", "regression testing", "security", "cybersecurity",
        "owasp", "siem", "soc", "network security", "vulnerability assessment",
        "computer science", "information technology", "cs / it", "programming",
        "data structures", "algorithms", "object-oriented", "oops",
    },
    "Analyst": {
        "data analyst", "business analyst", "research analyst", "data science", "machine learning",
        "artificial intelligence", "deep learning", "nlp", "computer vision", "python", "sql",
        "excel", "power bi", "tableau", "looker", "reporting", "dashboard", "kpi",
        "business intelligence", "forecasting", "statistics", "data visualization",
        "scikit-learn", "tensorflow", "pytorch", "pandas", "numpy", "spark", "etl",
        "feature engineering", "llm", "transformer", "market research", "requirements gathering",
        "stakeholder management", "process improvement", "financial modelling",
    },
    "Core_NonTech": {
        "embedded systems", "firmware", "hardware", "electronics", "electrical", "mechanical",
        "manufacturing", "automotive", "tractor", "industrial machinery", "microcontroller",
        "rtos", "pcb", "vhdl", "verilog", "uart", "spi", "i2c", "iot", "field service",
        "quality control", "operations", "supply chain", "procurement", "sales", "marketing",
        "human resources", "finance", "accounting", "administration", "business development",
        "recruitment", "customer service", "admission counsellor", "logistics", "retail",
        "civil", "maintenance", "utility", "transformers", "dg sets", "hvac",
    },
}

_DOMAIN_KEYS = list(DOMAINS.keys())
_DOMAIN_TEXTS = list(DOMAINS.values())
_cached_domain_embeddings: Optional[np.ndarray] = None

NON_JD_PATTERNS = [
    r"\bapplication guideline",
    r"\bsteps to apply\b",
    r"\bcareer portal\b",
    r"\bmanage your profile\b",
    r"\bcreate an account\b",
    r"\be-signing\b",
    r"\bcandidate email\b",
    r"\bfield specific instructions\b",
    r"\bnavigation and guidelines\b",
    r"\bbrowse the open role\b",
    r"\bhiring process\b",
]

JD_PATTERNS = [
    r"\bjob description\b",
    r"\bposition summary\b",
    r"\brole summary\b",
    r"\bresponsibilit",
    r"\bqualification",
    r"\brequirements?\b",
    r"\btechnical skills?\b",
    r"\bpreferred skills?\b",
    r"\bexperience\b",
    r"\bwhat you'll do\b",
    r"\babout the role\b",
]


# ── Domain classifier ────────────────────────────────────────────────────────

def _get_domain_embeddings(embedder: SentenceTransformer) -> np.ndarray:
    global _cached_domain_embeddings
    if _cached_domain_embeddings is None:
        _cached_domain_embeddings = embedder.encode(
            _DOMAIN_TEXTS, normalize_embeddings=True
        )
    return _cached_domain_embeddings


def _classify_semantic(
    text: str,
    embedder: SentenceTransformer,
) -> Tuple[str, float, Optional[str]]:
    jd_vec = embedder.encode(text, normalize_embeddings=True).reshape(1, -1)
    proto_vecs = _get_domain_embeddings(embedder)
    sims = cosine_similarity(jd_vec, proto_vecs)[0]

    sorted_idx = np.argsort(sims)[::-1]
    top_idx, second_idx = int(sorted_idx[0]), int(sorted_idx[1])
    top_score = float(sims[top_idx])
    second_score = float(sims[second_idx])

    secondary = _DOMAIN_KEYS[second_idx] if second_score >= max(0.15, top_score * 0.80) else None
    return _DOMAIN_KEYS[top_idx], top_score, secondary


def _classify_tfidf(text: str) -> Tuple[str, float, Optional[str]]:
    corpus = [text] + _DOMAIN_TEXTS
    tfidf = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    matrix = tfidf.fit_transform(corpus)
    sims = cosine_similarity(matrix[0], matrix[1:])[0]

    sorted_idx = np.argsort(sims)[::-1]
    top_idx, second_idx = int(sorted_idx[0]), int(sorted_idx[1])
    top_score = float(sims[top_idx])

    secondary = (
        _DOMAIN_KEYS[second_idx]
        if float(sims[second_idx]) > 0.05
        else None
    )
    return _DOMAIN_KEYS[top_idx], top_score, secondary


def normalize_text(text: str) -> str:
    text = text.lower()
    text = text.replace("node js", "node.js")
    text = text.replace("nodejs", "node.js")
    text = text.replace("react js", "react")
    text = text.replace("express js", "express")
    text = text.replace("full-stack", "full stack")
    return re.sub(r"\s+", " ", text).strip()


def extract_skills(text: str, term_bank: set[str]) -> set[str]:
    lowered = normalize_text(text)
    found = set()
    for term in term_bank:
        pattern = r"\b" + re.escape(term) + r"\b"
        if re.search(pattern, lowered):
            found.add(term)
    return found


def compute_domain_signal_counts(text: str) -> dict[str, int]:
    return {
        domain: len(extract_skills(text, DOMAIN_SKILL_TERMS[domain]))
        for domain in DOMAIN_SKILL_TERMS
    }


def has_software_role_signals(text: str) -> bool:
    lowered = normalize_text(text)
    role_patterns = [
        r"\bcs\s*/\s*it\b",
        r"\bcomputer science\b",
        r"\binformation technology\b",
        r"\bprogramming knowledge\b",
        r"\bdata structures\b",
        r"\balgorithms?\b",
        r"\bobject-oriented\b",
        r"\boops?\b",
        r"\bc/c\+\+\b",
        r"\bjava\b",
        r"\bpython\b",
    ]
    hits = sum(1 for pattern in role_patterns if re.search(pattern, lowered))
    return hits >= 3


def is_probable_jd(file_name: str, text: str) -> bool:
    combined = normalize_text(f"{file_name} {text}")
    has_jd_signals = any(re.search(pattern, combined) for pattern in JD_PATTERNS)
    has_non_jd_signals = any(re.search(pattern, combined) for pattern in NON_JD_PATTERNS)
    if has_non_jd_signals and not has_jd_signals:
        return False
    return has_jd_signals or not has_non_jd_signals


def classify_domain(
    text: str,
    embedder: Optional[SentenceTransformer] = None,
) -> Tuple[str, float, Optional[str]]:
    """
    Classify text into one of the predefined domains.

    Returns:
        (domain, confidence, secondary_domain)
        secondary_domain is None when the second-best score is too low.
    """
    if embedder is not None:
        semantic_domain, semantic_confidence, semantic_secondary = _classify_semantic(text, embedder)
    else:
        semantic_domain, semantic_confidence, semantic_secondary = _classify_tfidf(text)

    counts = compute_domain_signal_counts(text)
    combined_scores = {}
    software_signal_boost = has_software_role_signals(text)

    for domain in DOMAIN_SKILL_TERMS:
        term_component = min(1.0, counts[domain] / 6)
        semantic_component = semantic_confidence if domain == semantic_domain else 0.0
        score = (0.55 * semantic_component) + (0.45 * term_component)

        if software_signal_boost and domain == "Software":
            score += 0.18
        if software_signal_boost and domain == "Core_NonTech":
            score -= 0.10
        if domain == "Software" and counts["Software"] >= counts["Core_NonTech"] + 2:
            score += 0.08
        if domain == "Analyst" and counts["Analyst"] >= counts["Core_NonTech"] + 1:
            score += 0.06
        if domain == "Core_NonTech" and counts["Core_NonTech"] >= counts["Software"] + counts["Analyst"]:
            score += 0.06

        combined_scores[domain] = min(1.0, score)

    sorted_domains = sorted(combined_scores.items(), key=lambda item: item[1], reverse=True)
    best_domain, best_score = sorted_domains[0]
    second_domain, second_score = sorted_domains[1]

    secondary = None
    if second_score >= max(0.18, best_score * 0.80):
        secondary = second_domain
    elif semantic_secondary in DOMAIN_SKILL_TERMS and semantic_secondary != best_domain:
        secondary = semantic_secondary

    return best_domain, float(best_score), secondary


# ── Priority patterns for chunk relevance ────────────────────────────────────

PRIORITY_PATTERNS = [
    r"\brequired\b",
    r"\bpreferred\b",
    r"\bqualification",
    r"\bresponsibilit",
    r"\bjob description\b",
    r"\brole summary\b",
    r"\btechnical skills?\b",
    r"\bprogramming\b",
    r"\bsoftware\b",
    r"\bdeveloper\b",
    r"\bengineer\b",
    r"\bbackend\b",
    r"\bfrontend\b",
    r"\bfull[\s-]?stack\b",
    r"\bapi\b",
    r"\bdatabase\b",
    r"\bpython\b",
    r"\bjava(script)?\b",
    r"\breact\b",
    r"\bnode(\.js)?\b",
    r"\bc\+\+\b",
    r"\bsql\b",
    r"\baws\b",
    r"\bgit\b",
    r"\bdata\b",
    r"\bmachine learning\b",
    r"\bcloud\b",
    r"\bdevops\b",
    r"\bkubernetes\b",
    r"\bdocker\b",
    r"\bsecurity\b",
    r"\bembedded\b",
    r"\bfirmware\b",
]


# ── DB helpers ───────────────────────────────────────────────────────────────

def get_db_connection():
    if not DATABASE_URL:
        raise ValueError("DATABASE_URL is not set in environment")
    conn = psycopg2.connect(DATABASE_URL)
    cursor = conn.cursor()
    return conn, cursor


def ensure_table(cursor, conn) -> None:
    """
    Create job_descriptions table if it doesn't exist, then ensure
    all required columns (including domain fields) are present.
    Safe to run on a fresh or existing DB.
    """
    # Create table with all columns from scratch
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS job_descriptions (
            id                 SERIAL PRIMARY KEY,
            file_name          VARCHAR(255) UNIQUE NOT NULL,
            jd_text            TEXT,
            cleaned_text       TEXT,
            embedding          FLOAT[],
            domain             VARCHAR(60),
            domain_confidence  FLOAT,
            domain_secondary   VARCHAR(60),
            created_at         TIMESTAMP DEFAULT NOW()
        )
        """
    )
    conn.commit()

    # If the table already existed without domain columns, add them
    for col, col_type in [
        ("domain",            "VARCHAR(60)"),
        ("domain_confidence", "FLOAT"),
        ("domain_secondary",  "VARCHAR(60)"),
        ("created_at",        "TIMESTAMP DEFAULT NOW()"),
    ]:
        try:
            cursor.execute(
                f"ALTER TABLE job_descriptions ADD COLUMN IF NOT EXISTS {col} {col_type}"
            )
            conn.commit()
        except Exception:
            conn.rollback()

    # Index for fast domain-filtered queries
    try:
        cursor.execute(
            "CREATE INDEX IF NOT EXISTS idx_jd_domain ON job_descriptions(domain)"
        )
        conn.commit()
    except Exception:
        conn.rollback()

    print("  Table ready: job_descriptions")


def save_to_database(
    cursor,
    conn,
    file_name: str,
    jd_text: str,
    cleaned_text: str,
    embedding: List[float],
    domain: str,
    domain_confidence: float,
    domain_secondary: Optional[str],
) -> None:
    try:
        cursor.execute(
            """
            INSERT INTO job_descriptions
                (file_name, jd_text, cleaned_text, embedding,
                 domain, domain_confidence, domain_secondary)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (file_name) DO UPDATE
                SET jd_text            = EXCLUDED.jd_text,
                    cleaned_text       = EXCLUDED.cleaned_text,
                    embedding          = EXCLUDED.embedding,
                    domain             = EXCLUDED.domain,
                    domain_confidence  = EXCLUDED.domain_confidence,
                    domain_secondary   = EXCLUDED.domain_secondary
            """,
            (
                file_name, jd_text, cleaned_text,
                embedding if embedding else None,
                domain, domain_confidence, domain_secondary,
            ),
        )
        conn.commit()
        print(
            f"  Saved → domain: {domain} "
            f"(conf: {domain_confidence:.2f}"
            + (f", secondary: {domain_secondary}" if domain_secondary else "")
            + ")"
        )
    except Exception as e:
        conn.rollback()
        print(f"  DB insert failed for {file_name}: {e}")


# ── File reading ─────────────────────────────────────────────────────────────

def get_text_from_file(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        elif ext == ".docx" and docx is not None:
            document = docx.Document(path)
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            tables = [
                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(paragraphs + tables)
    except Exception as e:
        print(f"  Error reading {path}: {e}")
    return text.strip()


# ── Chunking ─────────────────────────────────────────────────────────────────

def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def split_into_logical_chunks(text: str) -> List[str]:
    parts = re.split(r"\n{2,}|(?=\s*[\u2022\-\*])", text)
    return [normalize_space(p) for p in parts if len(normalize_space(p)) >= 35]


def chunk_is_priority(chunk: str) -> bool:
    lowered = chunk.lower()
    return any(re.search(pat, lowered) for pat in PRIORITY_PATTERNS)


# ── Relevance filtering ───────────────────────────────────────────────────────

RELEVANCE_QUERIES = [
    "Required technical skills, programming languages, frameworks, tools, and education.",
    "Core job responsibilities, engineering tasks, deliverables, and day to day work.",
    "Candidate qualifications, must have requirements, preferred experience, and role expectations.",
]


def filter_relevant_chunks(
    chunks: List[str],
    embedder: Optional[SentenceTransformer],
) -> List[Tuple[str, float]]:
    """
    Score each chunk against relevance queries and return (chunk, score) pairs
    that pass the dynamic threshold, de-duplicated by near-exact similarity.
    """
    if not chunks:
        return []

    # ── Score chunks ──
    if embedder is not None:
        query_embs = embedder.encode(RELEVANCE_QUERIES, convert_to_tensor=True)
        chunk_embs = embedder.encode(chunks, convert_to_tensor=True)
        cos_scores = util.cos_sim(chunk_embs, query_embs)
        max_scores = cos_scores.max(dim=1).values.cpu().numpy()
    else:
        tfidf = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        matrix = tfidf.fit_transform(chunks + RELEVANCE_QUERIES)
        chunk_matrix = matrix[: len(chunks)]
        query_matrix = matrix[len(chunks):]
        cos_scores = cosine_similarity(chunk_matrix, query_matrix)
        max_scores = cos_scores.max(axis=1)

    mean_score = float(np.mean(max_scores))
    std_score = float(np.std(max_scores))
    dynamic_threshold = mean_score + (0.45 * std_score)

    # ── Apply priority boost and threshold ──
    scored: List[Tuple[str, float]] = []
    seen_texts: set = set()

    for i, raw_score in enumerate(max_scores):
        chunk = chunks[i]
        if chunk in seen_texts:
            continue

        boost = 0.0
        if chunk_is_priority(chunk):
            boost += 0.08
        if chunk.endswith(":") and len(chunk) < 50:
            boost += 0.08

        final = float(raw_score) + boost
        keep = final >= dynamic_threshold or (
            chunk_is_priority(chunk) and final >= mean_score
        )
        if keep:
            scored.append((chunk, round(final, 3)))
            seen_texts.add(chunk)

    if not scored:
        return []

    # ── De-duplicate near-identical chunks ──
    scored.sort(key=lambda x: x[1], reverse=True)
    texts = [t for t, _ in scored]

    if embedder is not None:
        embs = embedder.encode(texts, convert_to_tensor=True)
        sim_matrix = util.cos_sim(embs, embs).cpu().numpy()
    else:
        tfidf2 = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        embs2 = tfidf2.fit_transform(texts)
        sim_matrix = cosine_similarity(embs2, embs2)

    skipped: set = set()
    final_chunks: List[Tuple[str, float]] = []
    for i in range(len(texts)):
        if i in skipped:
            continue
        final_chunks.append(scored[i])
        for j in range(i + 1, len(texts)):
            if sim_matrix[i][j] > 0.95:
                skipped.add(j)

    return final_chunks


# ── Main pipeline ─────────────────────────────────────────────────────────────

def extract_and_store_jds(folder_path: str) -> None:
    """
    Process every PDF / DOCX in folder_path:
      1. Extract text
      2. Filter to relevant chunks
      3. Classify domain
      4. Save to DB with domain metadata
    """
    if not os.path.isdir(folder_path):
        print(f"Folder not found: {folder_path}")
        return

    conn, cursor = get_db_connection()
    print("DB connected successfully")

    # Create table if missing, add domain columns if upgrading an old table
    ensure_table(cursor, conn)

    embedder = get_model()
    if embedder is None:
        print("Using TF-IDF fallback (SentenceTransformer model not available locally)")

    files = [
        f for f in os.listdir(folder_path)
        if f.lower().endswith((".pdf", ".docx"))
    ]
    if not files:
        print(f"No PDF or DOCX files found in {folder_path}")
        return

    print(f"\nFound {len(files)} file(s) to process\n{'─' * 50}")

    try:
        for file_name in files:
            print(f"\nProcessing: {file_name}")
            file_path = os.path.join(folder_path, file_name)

            full_text = get_text_from_file(file_path)
            if not full_text:
                print("  No text extracted — skipping")
                continue

            if not is_probable_jd(file_name, full_text):
                print("  Looks like a non-JD document — skipping")
                continue

            title_hint = (
                os.path.splitext(file_name)[0]
                .replace("_", " ")
                .replace("-", " ")
            )
            chunks = [title_hint] + split_into_logical_chunks(full_text)

            relevant = filter_relevant_chunks(chunks, embedder)
            if not relevant:
                print("  No relevant content found — skipping")
                continue

            cleaned_text = "\n".join(text for text, _ in relevant)
            print(f"  Chunks kept: {len(relevant)} / {len(chunks)}")

            classification_input = f"{title_hint}\n{cleaned_text}"
            domain, confidence, secondary = classify_domain(
                classification_input, embedder=embedder
            )

            if embedder is not None:
                final_embedding: List[float] = embedder.encode(cleaned_text).tolist()
            else:
                final_embedding = []

            save_to_database(
                cursor=cursor,
                conn=conn,
                file_name=file_name,
                jd_text=full_text,
                cleaned_text=cleaned_text,
                embedding=final_embedding,
                domain=domain,
                domain_confidence=confidence,
                domain_secondary=secondary,
            )

    finally:
        cursor.close()
        conn.close()
        print(f"\n{'─' * 50}")
        print("All JDs processed. DB connection closed.")


# ── Domain stats (utility) ────────────────────────────────────────────────────

def print_domain_summary() -> None:
    """Print a count of JDs per domain stored in the DB."""
    conn, cursor = get_db_connection()
    try:
        cursor.execute(
            """
            SELECT domain, COUNT(*) AS count
            FROM job_descriptions
            WHERE domain IS NOT NULL
            GROUP BY domain
            ORDER BY count DESC
            """
        )
        rows = cursor.fetchall()
        print("\nDomain summary in DB:")
        print(f"  {'Domain':<25} {'Count':>6}")
        print(f"  {'─' * 32}")
        for domain, count in rows:
            print(f"  {domain:<25} {count:>6}")
    except Exception as e:
        print(f"Could not fetch domain summary: {e}")
    finally:
        cursor.close()
        conn.close()


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    JD_FOLDER = "./extracted/JDs"
    extract_and_store_jds(JD_FOLDER)
    print_domain_summary()
