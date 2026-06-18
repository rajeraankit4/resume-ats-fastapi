import os
import re
from typing import Dict, List, Optional, Set, Tuple

import fitz
import numpy as np
import psycopg2
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import docx
except ImportError:
    print("python-docx not found: pip install python-docx")
    docx = None


load_dotenv()
DATABASE_URL = os.getenv("DATABASE_URL")

DOMAINS = {
    "Software": "SDE, QA, DevOps, Cybersecurity, and software/platform roles",
    "Analyst": "Data Analyst, Business Analyst, ML/AI, and analytics roles",
    "Core_NonTech": "Core/NonTech roles",
}

DOMAIN_PROTOTYPES = {
    "Software": (
        "Software engineering backend frontend full stack QA test automation DevOps cloud "
        "infrastructure cybersecurity application development APIs web development databases."
    ),
    "Analyst": (
        "Data analyst business analyst research analyst data science machine learning "
        "artificial intelligence reporting dashboards business intelligence forecasting "
        "statistics data visualization experimentation."
    ),
    "Core_NonTech": (
        "Core engineering embedded systems firmware hardware electronics electrical mechanical "
        "manufacturing automotive operations procurement supply chain sales marketing HR finance "
        "administration non-technical roles."
    ),
}

DOMAIN_SKILL_TERMS: Dict[str, Set[str]] = {
    "Software": {
        "software", "developer", "backend", "frontend", "full stack", "web", "api", "rest api",
        "graphql", "node.js", "react", "angular", "vue", "express", "javascript", "typescript",
        "html", "css", "java", "python", "c++", "sql", "mongodb", "postgresql", "redis",
        "git", "github", "docker", "aws", "gcp", "azure", "linux", "kubernetes", "terraform",
        "jenkins", "github actions", "devops", "cloud", "microservices", "system design",
        "qa", "selenium", "cypress", "jest", "regression testing", "security", "cybersecurity",
        "owasp", "siem", "soc", "network security", "vulnerability assessment",
        "computer science", "information technology", "cs / it", "programming",
        "data structures", "algorithms", "object-oriented", "oops",
    },
    "Analyst": {
        "data analyst", "business analyst", "research analyst", "data science", "machine learning",
        "artificial intelligence", "deep learning", "nlp", "computer vision", "python", "sql",
        "excel", "power bi", "tableau", "looker", "reporting", "dashboard", "kpi",
        "business intelligence", "forecasting", "statistics", "data visualization",
        "scikit-learn", "tensorflow", "pytorch", "pandas", "numpy", "spark", "etl",
        "feature engineering", "llm", "transformer", "market research", "requirements gathering",
        "stakeholder management", "process improvement", "financial modelling",
    },
    "Core_NonTech": {
        "embedded systems", "firmware", "hardware", "electronics", "electrical", "mechanical",
        "manufacturing", "automotive", "tractor", "industrial machinery", "microcontroller",
        "rtos", "pcb", "vhdl", "verilog", "uart", "spi", "i2c", "iot", "field service",
        "quality control", "operations", "supply chain", "procurement", "sales", "marketing",
        "human resources", "finance", "accounting", "administration", "business development",
        "recruitment", "customer service", "admission counsellor", "logistics", "retail",
        "civil", "maintenance", "utility", "transformers", "dg sets", "hvac",
    },
}

DOMAIN_INTENT_PATTERNS = {
    "Software": [
        r"\bsoftware engineer",
        r"\bsoftware development\b",
        r"\bbackend\b",
        r"\bfrontend\b",
        r"\bfull stack\b",
        r"\bweb developer\b",
        r"\bapplication development\b",
        r"\bdevops\b",
        r"\bcybersecurity\b",
        r"\bqa\b",
    ],
    "Analyst": [
        r"\bdata science\b",
        r"\banalytics\b",
        r"\bdata analyst\b",
        r"\bbusiness analyst\b",
        r"\bmachine learning\b",
        r"\bartificial intelligence\b",
        r"\bpower bi\b",
        r"\bdashboard\b",
        r"\beda\b",
        r"\bdata analysis\b",
        r"\bdata-driven\b",
        r"\bsql queries?\b",
    ],
    "Core_NonTech": [
        r"\bembedded\b",
        r"\bfirmware\b",
        r"\belectrical\b",
        r"\bmechanical\b",
        r"\bmanufacturing\b",
        r"\bhardware\b",
        r"\boperations\b",
    ],
}

NON_JD_PATTERNS = [
    r"\bapplication guideline",
    r"\bsteps to apply\b",
    r"\bcareer portal\b",
    r"\bmanage your profile\b",
    r"\bcreate an account\b",
    r"\be-signing\b",
    r"\bcandidate email\b",
    r"\bfield specific instructions\b",
    r"\bnavigation and guidelines\b",
    r"\bbrowse the open role\b",
    r"\bhiring process\b",
]

JD_PATTERNS = [
    r"\bjob description\b",
    r"\bposition summary\b",
    r"\brole summary\b",
    r"\bresponsibilit",
    r"\bqualification",
    r"\brequirements?\b",
    r"\btechnical skills?\b",
    r"\bpreferred skills?\b",
    r"\bexperience\b",
    r"\bwhat you'll do\b",
    r"\babout the role\b",
]

SEMANTIC_RELEVANCE_GATE = 0.25
WEIGHT_SEMANTIC = 0.50
WEIGHT_SKILL = 0.30
WEIGHT_KEYWORD = 0.20
DOMAIN_FIT_WEIGHT = 0.65
DOMAIN_BUCKET_WEIGHT = 0.35

PRIORITY_PATTERNS = [
    r"\brequired\b",
    r"\bpreferred\b",
    r"\bqualification",
    r"\bresponsibilit",
    r"\bjob description\b",
    r"\brole summary\b",
    r"\btechnical skills?\b",
    r"\bprogramming\b",
    r"\bsoftware\b",
    r"\bdeveloper\b",
    r"\bengineer\b",
    r"\bbackend\b",
    r"\bfrontend\b",
    r"\bfull[\s-]?stack\b",
    r"\bapi\b",
    r"\bdatabase\b",
    r"\bpython\b",
    r"\bjava(script)?\b",
    r"\breact\b",
    r"\bnode(\.js)?\b",
    r"\bc\+\+\b",
    r"\bsql\b",
    r"\baws\b",
    r"\bgit\b",
    r"\bdata\b",
    r"\bmachine learning\b",
    r"\bcloud\b",
    r"\bdevops\b",
    r"\bkubernetes\b",
    r"\bdocker\b",
    r"\bsecurity\b",
    r"\bembedded\b",
    r"\bfirmware\b",
]

RELEVANCE_QUERIES = [
    "Required technical skills, programming languages, frameworks, tools, and education.",
    "Core job responsibilities, engineering tasks, deliverables, and day to day work.",
    "Candidate qualifications, must have requirements, preferred experience, and role expectations.",
]

_model = None
_model_load_attempted = False
_db_connection_failed = False
_cached_domain_embeddings: Optional[np.ndarray] = None


def get_model() -> Optional[SentenceTransformer]:
    global _model, _model_load_attempted
    if _model is not None:
        return _model
    if _model_load_attempted:
        return None
    _model_load_attempted = True
    try:
        _model = SentenceTransformer("all-MiniLM-L6-v2", local_files_only=True)
        return _model
    except Exception as e:
        print(f"SentenceTransformer unavailable, falling back to TF-IDF: {e}")
        return None


def get_db_connection():
    global _db_connection_failed
    if not DATABASE_URL or _db_connection_failed:
        return None, None
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cursor = conn.cursor()
        return conn, cursor
    except Exception as e:
        print(f"DB connection failed: {e}")
        _db_connection_failed = True
        return None, None


def normalize_text(text: str) -> str:
    text = text.lower()
    text = text.replace("node js", "node.js")
    text = text.replace("nodejs", "node.js")
    text = text.replace("react js", "react")
    text = text.replace("express js", "express")
    text = text.replace("full-stack", "full stack")
    return re.sub(r"\s+", " ", text).strip()


def _extract_docx_text(document) -> str:
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = [
        " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        for table in document.tables
        for row in table.rows
    ]
    return "\n".join(paragraphs + tables)


def get_text_from_file(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        elif ext == ".docx" and docx is not None:
            document = docx.Document(path)
            text = _extract_docx_text(document)
    except Exception as e:
        print(f"Error reading {path}: {e}")
    return text.strip()


def parse_embedding(raw) -> Optional[np.ndarray]:
    if raw is None:
        return None
    if isinstance(raw, list):
        arr = np.array(raw, dtype=float)
    else:
        arr = np.fromstring(str(raw).strip("[]"), sep=",", dtype=float)
    if arr.size == 0:
        return None
    return arr.reshape(1, -1)


def extract_skills(text: str, term_bank: Set[str]) -> Set[str]:
    lowered = normalize_text(text)
    found = set()
    for term in term_bank:
        if re.search(r"\b" + re.escape(term) + r"\b", lowered):
            found.add(term)
    return found


def compute_domain_signal_counts(text: str) -> dict[str, int]:
    return {
        domain: len(extract_skills(text, DOMAIN_SKILL_TERMS[domain]))
        for domain in DOMAIN_SKILL_TERMS
    }


def has_software_role_signals(text: str) -> bool:
    lowered = normalize_text(text)
    role_patterns = [
        r"\bcs\s*/\s*it\b",
        r"\bcomputer science\b",
        r"\binformation technology\b",
        r"\bprogramming knowledge\b",
        r"\bdata structures\b",
        r"\balgorithms?\b",
        r"\bobject-oriented\b",
        r"\boops?\b",
        r"\bc/c\+\+\b",
        r"\bjava\b",
        r"\bpython\b",
    ]
    hits = sum(1 for pattern in role_patterns if re.search(pattern, lowered))
    return hits >= 3


def is_probable_jd(file_name: str, text: str) -> bool:
    combined = normalize_text(f"{file_name} {text}")
    has_jd_signals = any(re.search(pattern, combined) for pattern in JD_PATTERNS)
    has_non_jd_signals = any(re.search(pattern, combined) for pattern in NON_JD_PATTERNS)
    if has_non_jd_signals and not has_jd_signals:
        return False
    return has_jd_signals or not has_non_jd_signals


def _get_domain_embeddings(embedder: SentenceTransformer) -> np.ndarray:
    global _cached_domain_embeddings
    if _cached_domain_embeddings is None:
        _cached_domain_embeddings = embedder.encode(
            list(DOMAIN_PROTOTYPES.values()), normalize_embeddings=True
        )
    return _cached_domain_embeddings


def _classify_semantic(
    text: str,
    embedder: SentenceTransformer,
) -> Tuple[str, float, Optional[str]]:
    domain_keys = list(DOMAIN_PROTOTYPES.keys())
    jd_vec = embedder.encode(text, normalize_embeddings=True).reshape(1, -1)
    proto_vecs = _get_domain_embeddings(embedder)
    sims = cosine_similarity(jd_vec, proto_vecs)[0]

    sorted_idx = np.argsort(sims)[::-1]
    top_idx, second_idx = int(sorted_idx[0]), int(sorted_idx[1])
    top_score = float(sims[top_idx])
    second_score = float(sims[second_idx])

    secondary = domain_keys[second_idx] if second_score >= max(0.15, top_score * 0.80) else None
    return domain_keys[top_idx], top_score, secondary


def _classify_tfidf(text: str) -> Tuple[str, float, Optional[str]]:
    domain_keys = list(DOMAIN_PROTOTYPES.keys())
    corpus = [text] + list(DOMAIN_PROTOTYPES.values())
    tfidf = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    matrix = tfidf.fit_transform(corpus)
    sims = cosine_similarity(matrix[0], matrix[1:])[0]

    sorted_idx = np.argsort(sims)[::-1]
    top_idx, second_idx = int(sorted_idx[0]), int(sorted_idx[1])
    top_score = float(sims[top_idx])

    secondary = domain_keys[second_idx] if float(sims[second_idx]) > 0.05 else None
    return domain_keys[top_idx], top_score, secondary


def classify_domain(
    text: str,
    embedder: Optional[SentenceTransformer] = None,
) -> Tuple[str, float, Optional[str]]:
    if embedder is not None:
        semantic_domain, semantic_confidence, semantic_secondary = _classify_semantic(text, embedder)
    else:
        semantic_domain, semantic_confidence, semantic_secondary = _classify_tfidf(text)

    counts = compute_domain_signal_counts(text)
    combined_scores = {}
    software_signal_boost = has_software_role_signals(text)

    for domain in DOMAIN_SKILL_TERMS:
        term_component = min(1.0, counts[domain] / 6)
        semantic_component = semantic_confidence if domain == semantic_domain else 0.0
        score = (0.55 * semantic_component) + (0.45 * term_component)

        if software_signal_boost and domain == "Software":
            score += 0.18
        if software_signal_boost and domain == "Core_NonTech":
            score -= 0.10

        if domain == "Software" and counts["Software"] >= counts["Core_NonTech"] + 2:
            score += 0.08
        if domain == "Analyst" and counts["Analyst"] >= counts["Core_NonTech"] + 1:
            score += 0.06
        if domain == "Core_NonTech" and counts["Core_NonTech"] >= counts["Software"] + counts["Analyst"]:
            score += 0.06

        combined_scores[domain] = min(1.0, score)

    sorted_domains = sorted(combined_scores.items(), key=lambda item: item[1], reverse=True)
    best_domain, best_score = sorted_domains[0]
    second_domain, second_score = sorted_domains[1]

    secondary = None
    if second_score >= max(0.18, best_score * 0.80):
        secondary = second_domain
    elif semantic_secondary in DOMAIN_SKILL_TERMS and semantic_secondary != best_domain:
        secondary = semantic_secondary

    return best_domain, float(best_score), secondary


def get_db_connection():
    if not DATABASE_URL:
        raise ValueError("DATABASE_URL is not set in environment")
    conn = psycopg2.connect(DATABASE_URL)
    cursor = conn.cursor()
    return conn, cursor


def ensure_table(cursor, conn) -> None:
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS job_descriptions (
            id                 SERIAL PRIMARY KEY,
            file_name          VARCHAR(255) UNIQUE NOT NULL,
            jd_text            TEXT,
            cleaned_text       TEXT,
            embedding          FLOAT[],
            domain             VARCHAR(60),
            domain_confidence  FLOAT,
            domain_secondary   VARCHAR(60),
            created_at         TIMESTAMP DEFAULT NOW()
        )
        """
    )
    conn.commit()

    for col, col_type in [
        ("domain", "VARCHAR(60)"),
        ("domain_confidence", "FLOAT"),
        ("domain_secondary", "VARCHAR(60)"),
        ("created_at", "TIMESTAMP DEFAULT NOW()"),
    ]:
        try:
            cursor.execute(
                f"ALTER TABLE job_descriptions ADD COLUMN IF NOT EXISTS {col} {col_type}"
            )
            conn.commit()
        except Exception:
            conn.rollback()

    try:
        cursor.execute(
            "CREATE INDEX IF NOT EXISTS idx_jd_domain ON job_descriptions(domain)"
        )
        conn.commit()
    except Exception:
        conn.rollback()


def save_to_database(
    cursor,
    conn,
    file_name: str,
    jd_text: str,
    cleaned_text: str,
    embedding: List[float],
    domain: str,
    domain_confidence: float,
    domain_secondary: Optional[str],
) -> None:
    try:
        cursor.execute(
            """
            INSERT INTO job_descriptions
                (file_name, jd_text, cleaned_text, embedding,
                 domain, domain_confidence, domain_secondary)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (file_name) DO UPDATE
                SET jd_text            = EXCLUDED.jd_text,
                    cleaned_text       = EXCLUDED.cleaned_text,
                    embedding          = EXCLUDED.embedding,
                    domain             = EXCLUDED.domain,
                    domain_confidence  = EXCLUDED.domain_confidence,
                    domain_secondary   = EXCLUDED.domain_secondary
            """,
            (
                file_name,
                jd_text,
                cleaned_text,
                embedding if embedding else None,
                domain,
                domain_confidence,
                domain_secondary,
            ),
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        print(f"DB insert failed for {file_name}: {e}")


def get_text_from_file(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        elif ext == ".docx" and docx is not None:
            document = docx.Document(path)
            text = _extract_docx_text(document)
    except Exception as e:
        print(f"Error reading {path}: {e}")
    return text.strip()


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def split_into_logical_chunks(text: str) -> List[str]:
    parts = re.split(r"\n{2,}|(?=\s*[\u2022\-\*])", text)
    return [normalize_space(p) for p in parts if len(normalize_space(p)) >= 35]


def chunk_is_priority(chunk: str) -> bool:
    lowered = chunk.lower()
    return any(re.search(pat, lowered) for pat in PRIORITY_PATTERNS)


def filter_relevant_chunks(
    chunks: List[str],
    embedder: Optional[SentenceTransformer],
) -> List[Tuple[str, float]]:
    if not chunks:
        return []

    if embedder is not None:
        query_embs = embedder.encode(RELEVANCE_QUERIES, convert_to_tensor=True)
        chunk_embs = embedder.encode(chunks, convert_to_tensor=True)
        cos_scores = util.cos_sim(chunk_embs, query_embs)
        max_scores = cos_scores.max(dim=1).values.cpu().numpy()
    else:
        tfidf = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        matrix = tfidf.fit_transform(chunks + RELEVANCE_QUERIES)
        chunk_matrix = matrix[: len(chunks)]
        query_matrix = matrix[len(chunks):]
        cos_scores = cosine_similarity(chunk_matrix, query_matrix)
        max_scores = cos_scores.max(axis=1)

    mean_score = float(np.mean(max_scores))
    std_score = float(np.std(max_scores))
    dynamic_threshold = mean_score + (0.45 * std_score)

    scored: List[Tuple[str, float]] = []
    seen_texts: set = set()

    for i, raw_score in enumerate(max_scores):
        chunk = chunks[i]
        if chunk in seen_texts:
            continue

        boost = 0.0
        if chunk_is_priority(chunk):
            boost += 0.08
        if chunk.endswith(":") and len(chunk) < 50:
            boost += 0.08

        final = float(raw_score) + boost
        keep = final >= dynamic_threshold or (chunk_is_priority(chunk) and final >= mean_score)
        if keep:
            scored.append((chunk, round(final, 3)))
            seen_texts.add(chunk)

    if not scored:
        return []

    scored.sort(key=lambda x: x[1], reverse=True)
    texts = [t for t, _ in scored]

    if embedder is not None:
        embs = embedder.encode(texts, convert_to_tensor=True)
        sim_matrix = util.cos_sim(embs, embs).cpu().numpy()
    else:
        tfidf2 = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        embs2 = tfidf2.fit_transform(texts)
        sim_matrix = cosine_similarity(embs2, embs2)

    skipped: set = set()
    final_chunks: List[Tuple[str, float]] = []
    for i in range(len(texts)):
        if i in skipped:
            continue
        final_chunks.append(scored[i])
        for j in range(i + 1, len(texts)):
            if sim_matrix[i][j] > 0.95:
                skipped.add(j)

    return final_chunks


def fetch_jds_by_domain(domain: str) -> List[Dict]:
    conn, cursor = get_db_connection()
    if not conn:
        return []
    try:
        cursor.execute(
            """
            SELECT file_name, jd_text, cleaned_text, embedding,
                   domain, domain_confidence, domain_secondary
            FROM job_descriptions
            WHERE domain = %s OR domain_secondary = %s
            ORDER BY
                CASE WHEN domain = %s THEN 0 ELSE 1 END,
                domain_confidence DESC
            """,
            (domain, domain, domain),
        )
        rows = cursor.fetchall()
        return [
            {
                "file_name": row[0],
                "jd_text": row[1] or "",
                "cleaned_text": row[2] or row[1] or "",
                "embedding": row[3],
                "domain": row[4],
                "domain_confidence": row[5],
                "domain_secondary": row[6],
            }
            for row in rows
        ]
    except Exception as e:
        print(f"DB fetch error: {e}")
        return []
    finally:
        cursor.close()
        conn.close()


def fetch_jds_from_folder(folder_path: str) -> List[Dict]:
    if not os.path.isdir(folder_path):
        return []
    records = []
    for file_name in os.listdir(folder_path):
        if file_name.startswith("~$"):
            continue
        if not file_name.lower().endswith((".pdf", ".docx")):
            continue
        jd_text = get_text_from_file(os.path.join(folder_path, file_name))
        if not jd_text or not is_probable_jd(file_name, jd_text):
            continue
        title_hint = os.path.splitext(file_name)[0].replace("_", " ").replace("-", " ")
        records.append(
            {
                "file_name": file_name,
                "jd_text": jd_text,
                "cleaned_text": f"{title_hint}\n{jd_text}",
                "embedding": None,
                "domain": None,
                "domain_confidence": None,
                "domain_secondary": None,
            }
        )
    return records


def compute_semantic_scores(
    resume_text: str,
    jds: List[Dict],
    embedder: Optional[SentenceTransformer],
) -> List[float]:
    if embedder is not None:
        resume_vec = embedder.encode(resume_text, normalize_embeddings=True).reshape(1, -1)
        scores = []
        for jd in jds:
            jd_vec = parse_embedding(jd.get("embedding"))
            if jd_vec is None:
                jd_vec = embedder.encode(jd["cleaned_text"], normalize_embeddings=True).reshape(1, -1)
            scores.append(float(cosine_similarity(resume_vec, jd_vec)[0][0]))
        return scores

    corpus = [resume_text] + [jd["cleaned_text"] or jd["jd_text"] or jd["file_name"] for jd in jds]
    matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2)).fit_transform(corpus)
    resume_vec = matrix[0]
    return [float(cosine_similarity(resume_vec, matrix[i + 1])[0][0]) for i in range(len(jds))]


def score_skill_overlap(resume_skills: Set[str], jd_skills: Set[str]) -> float:
    if not jd_skills:
        return 0.0
    overlap = resume_skills & jd_skills
    precision = len(overlap) / len(jd_skills)
    recall = len(overlap) / max(len(resume_skills), 1)
    return min(1.0, 0.7 * precision + 0.3 * recall)


def score_keyword_density(resume_text: str, jd_text: str) -> float:
    stopwords = {
        "with", "that", "this", "have", "will", "from", "they", "been", "your", "team",
        "work", "role", "able", "good", "also", "must", "strong", "well", "over", "into",
        "should", "would", "could", "their", "about", "which", "after", "other", "some", "such",
    }
    jd_words = {w for w in re.findall(r"\b[a-z]{4,}\b", jd_text.lower()) if w not in stopwords}
    if not jd_words:
        return 0.0
    resume_lower = resume_text.lower()
    hits = sum(1 for w in jd_words if re.search(r"\b" + re.escape(w) + r"\b", resume_lower))
    return min(1.0, hits / len(jd_words))


def compute_resume_domain_fit(
    resume_text: str,
    domain: str,
    embedder: Optional[SentenceTransformer],
) -> float:
    normalized_resume = normalize_text(resume_text)
    domain_skills = extract_skills(resume_text, DOMAIN_SKILL_TERMS[domain])
    target_component = min(1.0, len(domain_skills) / 6)
    core_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Core_NonTech"]))
    software_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Software"]))
    analyst_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Analyst"]))
    intent_hits = sum(1 for pattern in DOMAIN_INTENT_PATTERNS[domain] if re.search(pattern, normalized_resume))
    intent_component = min(1.0, intent_hits / 4)

    if embedder is not None:
        resume_vec = embedder.encode(normalized_resume, normalize_embeddings=True).reshape(1, -1)
        domain_vec = embedder.encode(DOMAIN_PROTOTYPES[domain], normalize_embeddings=True).reshape(1, -1)
        semantic_component = float(cosine_similarity(resume_vec, domain_vec)[0][0])
    else:
        matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2)).fit_transform(
            [normalized_resume, DOMAIN_PROTOTYPES[domain]]
        )
        semantic_component = float(cosine_similarity(matrix[0], matrix[1])[0][0])

    fit = (0.45 * target_component) + (0.25 * semantic_component) + (0.30 * intent_component)

    if domain == "Core_NonTech" and software_component >= 6 and analyst_component <= 2 and core_component <= 2:
        fit *= 0.20
    if domain == "Software" and software_component >= analyst_component + core_component:
        fit += 0.08
    if domain == "Software" and analyst_component >= 5 and analyst_component >= core_component + 3:
        fit *= 0.88
    if domain == "Analyst" and analyst_component >= core_component:
        fit += 0.04
    if domain == "Analyst" and analyst_component >= 5 and intent_hits >= 2:
        fit += 0.10

    return max(0.0, min(1.0, fit))


def jd_matches_domain(domain: str, jd: Dict) -> bool:
    stored_domain = jd.get("domain")
    stored_conf = jd.get("domain_confidence") or 0.0
    if stored_domain == domain and stored_conf >= 0.55:
        return True
    if jd.get("domain_secondary") == domain and stored_conf >= 0.35:
        return True

    combined = f"{jd['file_name']}\n{jd['cleaned_text'] or jd['jd_text']}"
    counts = {name: len(extract_skills(combined, DOMAIN_SKILL_TERMS[name])) for name in DOMAIN_SKILL_TERMS}
    best_domain = max(counts, key=counts.get)

    if domain == "Software":
        return counts["Software"] >= 2 or (best_domain == "Software" and counts["Software"] >= 1)
    if domain == "Analyst":
        return counts["Analyst"] >= 2 or (best_domain == "Analyst" and counts["Analyst"] >= 1)
    return counts["Core_NonTech"] >= 2 or (
        best_domain == "Core_NonTech" and counts["Core_NonTech"] >= max(counts["Software"], counts["Analyst"])
    )


def compute_ats_score(semantic: float, skill_overlap: float, keyword_density: float) -> float:
    raw = (
        WEIGHT_SEMANTIC * semantic
        + WEIGHT_SKILL * skill_overlap
        + WEIGHT_KEYWORD * keyword_density
    )
    return max(0.0, min(1.0, raw))


def build_reason(overlap: Set[str], semantic: float, skill_score: float, keyword_score: float) -> str:
    parts = []
    if overlap:
        parts.append(f"matched skills: {', '.join(sorted(overlap)[:6])}")
    parts.append(
        "strong semantic match"
        if semantic >= 0.65
        else "moderate semantic match"
        if semantic >= 0.45
        else "low semantic similarity"
    )
    if skill_score < 0.20:
        parts.append("few domain skills in resume")
    if keyword_score < 0.20:
        parts.append("low JD keyword coverage")
    return "; ".join(parts)


def score_resume_against_jds(
    resume_text: str,
    domain: str,
    jds: List[Dict],
    top_n: int = 10,
    embedder: Optional[SentenceTransformer] = None,
) -> List[Dict]:
    if domain not in DOMAINS:
        return []
    if not resume_text or not jds:
        return []

    embedder = embedder if embedder is not None else get_model()
    skill_terms = DOMAIN_SKILL_TERMS[domain]
    resume_skills = extract_skills(resume_text, skill_terms)
    semantic_scores = compute_semantic_scores(resume_text, jds, embedder)
    results = []

    for jd, semantic in zip(jds, semantic_scores):
        if not jd_matches_domain(domain, jd):
            continue
        if semantic < SEMANTIC_RELEVANCE_GATE:
            continue

        jd_text = jd["jd_text"]
        cleaned = jd["cleaned_text"] or jd_text
        jd_skills = extract_skills(f"{jd['file_name']}\n{cleaned}\n{jd_text}", skill_terms)
        skill_overlap = score_skill_overlap(resume_skills, jd_skills)
        keyword_density = score_keyword_density(resume_text, jd_text)
        ats = compute_ats_score(semantic, skill_overlap, keyword_density)
        overlap_terms = resume_skills & jd_skills

        results.append(
            {
                "file_name": jd["file_name"],
                "ats_score": round(ats * 100, 1),
                "semantic": round(semantic * 100, 1),
                "skill_overlap": round(skill_overlap * 100, 1),
                "keyword_density": round(keyword_density * 100, 1),
                "matched_skills": sorted(overlap_terms),
                "reason": build_reason(overlap_terms, semantic, skill_overlap, keyword_density),
            }
        )

    results.sort(key=lambda item: item["ats_score"], reverse=True)
    return results[:top_n]


def match_resume(
    resume_path: str,
    domain: str,
    top_n: int = 10,
    jd_folder: str = "./extracted/JDs",
    verbose: bool = True,
) -> List[Dict]:
    if domain not in DOMAINS:
        if verbose:
            print(f"Unknown domain '{domain}'. Choose from: {list(DOMAINS.keys())}")
        return []

    resume_text = get_text_from_file(resume_path)
    if not resume_text:
        if verbose:
            print("ERROR: Could not extract text from resume.")
        return []

    jds = fetch_jds_by_domain(domain)
    if not jds:
        jds = fetch_jds_from_folder(jd_folder)

    return score_resume_against_jds(resume_text, domain, jds, top_n=top_n)


def match_resume_all_domains(resume_path: str, top_n_per_domain: int = 3) -> List[Dict]:
    resume_text = get_text_from_file(resume_path)
    embedder = get_model()
    if not resume_text:
        return []

    summary_rows = []
    for domain, label in DOMAINS.items():
        results = match_resume(resume_path, domain, top_n=top_n_per_domain, verbose=False)
        domain_fit = round(compute_resume_domain_fit(resume_text, domain, embedder) * 100, 1)

        if results:
            top_score = results[0]["ats_score"]
            top_file = results[0]["file_name"]
            top_matches = results[: min(3, len(results))]
            avg_top = round(sum(item["ats_score"] for item in top_matches) / len(top_matches), 1)
            hits = len(results)
        else:
            top_score = 0.0
            avg_top = 0.0
            top_file = "-"
            hits = 0

        overall_score = round((DOMAIN_FIT_WEIGHT * domain_fit) + (DOMAIN_BUCKET_WEIGHT * avg_top), 1)
        summary_rows.append(
            {
                "domain": domain,
                "label": label,
                "domain_fit": domain_fit,
                "top_score": top_score,
                "avg_top": avg_top,
                "overall_score": overall_score,
                "hits": hits,
                "top_file": top_file,
            }
        )

    summary_rows.sort(key=lambda item: (item["overall_score"], item["domain_fit"], item["avg_top"]), reverse=True)
    return summary_rows


def match_resume_against_jds(
    resume_path: str,
    jds: List[Dict],
    domain: str,
    top_n: int = 10,
) -> List[Dict]:
    resume_text = get_text_from_file(resume_path)
    return score_resume_against_jds(resume_text, domain, jds, top_n=top_n)


def match_resume_all_domains_against_jds(
    resume_path: str,
    jds: List[Dict],
    top_n_per_domain: int = 3,
) -> List[Dict]:
    resume_text = get_text_from_file(resume_path)
    embedder = get_model()
    if not resume_text:
        return []

    summary_rows = []
    for domain, label in DOMAINS.items():
        results = score_resume_against_jds(resume_text, domain, jds, top_n=top_n_per_domain)
        domain_fit = round(compute_resume_domain_fit(resume_text, domain, embedder) * 100, 1)

        if results:
            top_score = results[0]["ats_score"]
            top_file = results[0]["file_name"]
            top_matches = results[: min(3, len(results))]
            avg_top = round(sum(item["ats_score"] for item in top_matches) / len(top_matches), 1)
            hits = len(results)
        else:
            top_score = 0.0
            avg_top = 0.0
            top_file = "-"
            hits = 0

        overall_score = round((DOMAIN_FIT_WEIGHT * domain_fit) + (DOMAIN_BUCKET_WEIGHT * avg_top), 1)
        summary_rows.append(
            {
                "domain": domain,
                "label": label,
                "domain_fit": domain_fit,
                "top_score": top_score,
                "avg_top": avg_top,
                "overall_score": overall_score,
                "hits": hits,
                "top_file": top_file,
            }
        )

    summary_rows.sort(key=lambda item: (item["overall_score"], item["domain_fit"], item["avg_top"]), reverse=True)
    return summary_rows


def extract_and_store_jds(folder_path: str) -> None:
    if not os.path.isdir(folder_path):
        print(f"Folder not found: {folder_path}")
        return

    conn, cursor = get_db_connection()
    ensure_table(cursor, conn)

    embedder = get_model()
    files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".docx"))]
    if not files:
        print(f"No PDF or DOCX files found in {folder_path}")
        return

    try:
        for file_name in files:
            file_path = os.path.join(folder_path, file_name)
            full_text = get_text_from_file(file_path)
            if not full_text or not is_probable_jd(file_name, full_text):
                continue

            title_hint = os.path.splitext(file_name)[0].replace("_", " ").replace("-", " ")
            chunks = [title_hint] + split_into_logical_chunks(full_text)
            relevant = filter_relevant_chunks(chunks, embedder)
            if not relevant:
                continue

            cleaned_text = "\n".join(text for text, _ in relevant)
            classification_input = f"{title_hint}\n{cleaned_text}"
            domain, confidence, secondary = classify_domain(classification_input, embedder=embedder)

            if embedder is not None:
                final_embedding: List[float] = embedder.encode(cleaned_text).tolist()
            else:
                final_embedding = []

            save_to_database(
                cursor=cursor,
                conn=conn,
                file_name=file_name,
                jd_text=full_text,
                cleaned_text=cleaned_text,
                embedding=final_embedding,
                domain=domain,
                domain_confidence=confidence,
                domain_secondary=secondary,
            )
    finally:
        cursor.close()
        conn.close()


def print_domain_summary() -> None:
    conn, cursor = get_db_connection()
    try:
        cursor.execute(
            """
            SELECT domain, COUNT(*) AS count
            FROM job_descriptions
            WHERE domain IS NOT NULL
            GROUP BY domain
            ORDER BY count DESC
            """
        )
        rows = cursor.fetchall()
        print("\nDomain summary in DB:")
        print(f"  {'Domain':<25} {'Count':>6}")
        print(f"  {'─' * 32}")
        for domain, count in rows:
            print(f"  {domain:<25} {count:>6}")
    except Exception as e:
        print(f"Could not fetch domain summary: {e}")
    finally:
        cursor.close()
        conn.close()
