import os
import re
from typing import Optional

import fitz
import numpy as np

try:
    import docx
except ImportError:
    print("python-docx not found: pip install python-docx")
    docx = None


def normalize_text(text: str) -> str:
    text = text.lower()
    text = text.replace("node js", "node.js")
    text = text.replace("nodejs", "node.js")
    text = text.replace("react js", "react")
    text = text.replace("express js", "express")
    text = text.replace("full-stack", "full stack")
    return re.sub(r"\s+", " ", text).strip()


def get_text_from_file(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        elif ext == ".docx" and docx is not None:
            document = docx.Document(path)
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            tables = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(paragraphs + tables)
    except Exception as e:
        print(f"Error reading {path}: {e}")
    return text.strip()


def parse_embedding(raw) -> Optional[np.ndarray]:
    if raw is None:
        return None
    if isinstance(raw, list):
        arr = np.array(raw, dtype=float)
    else:
        arr = np.fromstring(str(raw).strip("[]"), sep=",", dtype=float)
    if arr.size == 0:
        return None
    return arr.reshape(1, -1)

