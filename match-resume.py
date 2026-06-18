import os
import re
import sys
from typing import Dict, List, Optional, Set, Tuple

import fitz
import numpy as np
import psycopg2
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import docx
except ImportError:
    print("python-docx not found: pip install python-docx")
    docx = None


load_dotenv()
DATABASE_URL = os.getenv("DATABASE_URL")

DOMAINS = {
    "Software": "SDE, QA, DevOps, Cybersecurity, and software/platform roles",
    "Analyst": "Data Analyst, Business Analyst, ML/AI, and analytics roles",
    "Core_NonTech": "Core/NonTech roles",
}

DOMAIN_PROTOTYPES = {
    "Software": (
        "Software engineering backend frontend full stack QA test automation DevOps cloud "
        "infrastructure cybersecurity application development APIs web development databases."
    ),
    "Analyst": (
        "Data analyst business analyst research analyst data science machine learning "
        "artificial intelligence reporting dashboards business intelligence forecasting "
        "statistics data visualization experimentation."
    ),
    "Core_NonTech": (
        "Core engineering embedded systems firmware hardware electronics electrical mechanical "
        "manufacturing automotive operations procurement supply chain sales marketing HR finance "
        "administration non-technical roles."
    ),
}

DOMAIN_SKILL_TERMS: Dict[str, Set[str]] = {
    "Software": {
        "software", "developer", "backend", "frontend", "full stack", "web", "api", "rest api",
        "graphql", "node.js", "react", "angular", "vue", "express", "javascript", "typescript",
        "html", "css", "java", "python", "c++", "sql", "mongodb", "postgresql", "redis",
        "git", "github", "docker", "aws", "gcp", "azure", "linux", "kubernetes", "terraform",
        "jenkins", "github actions", "devops", "cloud", "microservices", "system design",
        "qa", "selenium", "cypress", "jest", "regression testing", "security", "cybersecurity",
        "owasp", "siem", "soc", "network security", "vulnerability assessment",
        "computer science", "information technology", "cs / it", "programming",
        "data structures", "algorithms", "object-oriented", "oops",
    },
    "Analyst": {
        "data analyst", "business analyst", "research analyst", "data science", "machine learning",
        "artificial intelligence", "deep learning", "nlp", "computer vision", "python", "sql",
        "excel", "power bi", "tableau", "looker", "reporting", "dashboard", "kpi",
        "business intelligence", "forecasting", "statistics", "data visualization",
        "scikit-learn", "tensorflow", "pytorch", "pandas", "numpy", "spark", "etl",
        "feature engineering", "llm", "transformer", "market research", "requirements gathering",
        "stakeholder management", "process improvement", "financial modelling",
    },
    "Core_NonTech": {
        "embedded systems", "firmware", "hardware", "electronics", "electrical", "mechanical",
        "manufacturing", "automotive", "tractor", "industrial machinery", "microcontroller",
        "rtos", "pcb", "vhdl", "verilog", "uart", "spi", "i2c", "iot", "field service",
        "quality control", "operations", "supply chain", "procurement", "sales", "marketing",
        "human resources", "finance", "accounting", "administration", "business development",
        "recruitment", "customer service", "admission counsellor", "logistics", "retail",
        "civil", "maintenance", "utility", "transformers", "dg sets", "hvac",
    },
}

DOMAIN_INTENT_PATTERNS = {
    "Software": [
        r"\bsoftware engineer",
        r"\bsoftware development\b",
        r"\bbackend\b",
        r"\bfrontend\b",
        r"\bfull stack\b",
        r"\bweb developer\b",
        r"\bapplication development\b",
        r"\bdevops\b",
        r"\bcybersecurity\b",
        r"\bqa\b",
    ],
    "Analyst": [
        r"\bdata science\b",
        r"\banalytics\b",
        r"\bdata analyst\b",
        r"\bbusiness analyst\b",
        r"\bmachine learning\b",
        r"\bartificial intelligence\b",
        r"\bpower bi\b",
        r"\bdashboard\b",
        r"\beda\b",
        r"\bdata analysis\b",
        r"\bdata-driven\b",
        r"\bsql queries?\b",
    ],
    "Core_NonTech": [
        r"\bembedded\b",
        r"\bfirmware\b",
        r"\belectrical\b",
        r"\bmechanical\b",
        r"\bmanufacturing\b",
        r"\bhardware\b",
        r"\boperations\b",
    ],
}

NON_JD_PATTERNS = [
    r"\bapplication guideline",
    r"\bsteps to apply\b",
    r"\bcareer portal\b",
    r"\bmanage your profile\b",
    r"\bcreate an account\b",
    r"\be-signing\b",
    r"\bcandidate email\b",
    r"\bfield specific instructions\b",
    r"\bnavigation and guidelines\b",
    r"\bbrowse the open role\b",
    r"\bhiring process\b",
]

JD_PATTERNS = [
    r"\bjob description\b",
    r"\bposition summary\b",
    r"\brole summary\b",
    r"\bresponsibilit",
    r"\bqualification",
    r"\brequirements?\b",
    r"\btechnical skills?\b",
    r"\bpreferred skills?\b",
    r"\bexperience\b",
    r"\bwhat you'll do\b",
    r"\babout the role\b",
]

SEMANTIC_RELEVANCE_GATE = 0.25
WEIGHT_SEMANTIC = 0.50
WEIGHT_SKILL = 0.30
WEIGHT_KEYWORD = 0.20
DOMAIN_FIT_WEIGHT = 0.65
DOMAIN_BUCKET_WEIGHT = 0.35

_model = None
_model_load_attempted = False
_db_connection_failed = False


def get_model() -> Optional[SentenceTransformer]:
    global _model, _model_load_attempted
    if _model is not None:
        return _model
    if _model_load_attempted:
        return None
    _model_load_attempted = True
    try:
        _model = SentenceTransformer("all-MiniLM-L6-v2", local_files_only=True)
        return _model
    except Exception as e:
        print(f"SentenceTransformer unavailable, falling back to TF-IDF: {e}")
        return None


def get_db_connection():
    global _db_connection_failed
    if not DATABASE_URL or _db_connection_failed:
        return None, None
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cursor = conn.cursor()
        return conn, cursor
    except Exception as e:
        print(f"DB connection failed: {e}")
        _db_connection_failed = True
        return None, None


def normalize_text(text: str) -> str:
    text = text.lower()
    text = text.replace("node js", "node.js")
    text = text.replace("nodejs", "node.js")
    text = text.replace("react js", "react")
    text = text.replace("express js", "express")
    text = text.replace("full-stack", "full stack")
    return re.sub(r"\s+", " ", text).strip()


def get_text_from_file(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        elif ext == ".docx" and docx is not None:
            document = docx.Document(path)
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            tables = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(paragraphs + tables)
    except Exception as e:
        print(f"Error reading {path}: {e}")
    return text.strip()


def parse_embedding(raw) -> Optional[np.ndarray]:
    if raw is None:
        return None
    if isinstance(raw, list):
        arr = np.array(raw, dtype=float)
    else:
        arr = np.fromstring(str(raw).strip("[]"), sep=",", dtype=float)
    if arr.size == 0:
        return None
    return arr.reshape(1, -1)


def extract_skills(text: str, term_bank: Set[str]) -> Set[str]:
    lowered = normalize_text(text)
    found = set()
    for term in term_bank:
        if re.search(r"\b" + re.escape(term) + r"\b", lowered):
            found.add(term)
    return found


def is_probable_jd(file_name: str, text: str) -> bool:
    combined = normalize_text(f"{file_name} {text}")
    has_jd_signals = any(re.search(pattern, combined) for pattern in JD_PATTERNS)
    has_non_jd_signals = any(re.search(pattern, combined) for pattern in NON_JD_PATTERNS)
    if has_non_jd_signals and not has_jd_signals:
        return False
    return has_jd_signals or not has_non_jd_signals


def fetch_jds_by_domain(domain: str) -> List[Dict]:
    conn, cursor = get_db_connection()
    if not conn:
        return []
    try:
        cursor.execute(
            """
            SELECT file_name, jd_text, cleaned_text, embedding,
                   domain, domain_confidence, domain_secondary
            FROM job_descriptions
            WHERE domain = %s OR domain_secondary = %s
            ORDER BY
                CASE WHEN domain = %s THEN 0 ELSE 1 END,
                domain_confidence DESC
            """,
            (domain, domain, domain),
        )
        rows = cursor.fetchall()
        return [
            {
                "file_name": row[0],
                "jd_text": row[1] or "",
                "cleaned_text": row[2] or row[1] or "",
                "embedding": row[3],
                "domain": row[4],
                "domain_confidence": row[5],
                "domain_secondary": row[6],
            }
            for row in rows
        ]
    except Exception as e:
        print(f"DB fetch error: {e}")
        return []
    finally:
        cursor.close()
        conn.close()


def fetch_jds_from_folder(folder_path: str) -> List[Dict]:
    if not os.path.isdir(folder_path):
        return []
    records = []
    for file_name in os.listdir(folder_path):
        if file_name.startswith("~$"):
            continue
        if not file_name.lower().endswith((".pdf", ".docx")):
            continue
        jd_text = get_text_from_file(os.path.join(folder_path, file_name))
        if not jd_text or not is_probable_jd(file_name, jd_text):
            continue
        title_hint = os.path.splitext(file_name)[0].replace("_", " ").replace("-", " ")
        records.append(
            {
                "file_name": file_name,
                "jd_text": jd_text,
                "cleaned_text": f"{title_hint}\n{jd_text}",
                "embedding": None,
                "domain": None,
                "domain_confidence": None,
                "domain_secondary": None,
            }
        )
    return records


def compute_semantic_scores(
    resume_text: str,
    jds: List[Dict],
    embedder: Optional[SentenceTransformer],
) -> List[float]:
    if embedder is not None:
        resume_vec = embedder.encode(resume_text, normalize_embeddings=True).reshape(1, -1)
        scores = []
        for jd in jds:
            jd_vec = parse_embedding(jd.get("embedding"))
            if jd_vec is None:
                jd_vec = embedder.encode(jd["cleaned_text"], normalize_embeddings=True).reshape(1, -1)
            scores.append(float(cosine_similarity(resume_vec, jd_vec)[0][0]))
        return scores

    corpus = [resume_text] + [jd["cleaned_text"] or jd["jd_text"] or jd["file_name"] for jd in jds]
    matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2)).fit_transform(corpus)
    resume_vec = matrix[0]
    return [float(cosine_similarity(resume_vec, matrix[i + 1])[0][0]) for i in range(len(jds))]


def score_skill_overlap(resume_skills: Set[str], jd_skills: Set[str]) -> float:
    if not jd_skills:
        return 0.0
    overlap = resume_skills & jd_skills
    precision = len(overlap) / len(jd_skills)
    recall = len(overlap) / max(len(resume_skills), 1)
    return min(1.0, 0.7 * precision + 0.3 * recall)


def score_keyword_density(resume_text: str, jd_text: str) -> float:
    stopwords = {
        "with", "that", "this", "have", "will", "from", "they", "been", "your", "team",
        "work", "role", "able", "good", "also", "must", "strong", "well", "over", "into",
        "should", "would", "could", "their", "about", "which", "after", "other", "some", "such",
    }
    jd_words = {w for w in re.findall(r"\b[a-z]{4,}\b", jd_text.lower()) if w not in stopwords}
    if not jd_words:
        return 0.0
    resume_lower = resume_text.lower()
    hits = sum(1 for w in jd_words if re.search(r"\b" + re.escape(w) + r"\b", resume_lower))
    return min(1.0, hits / len(jd_words))


def compute_resume_domain_fit(
    resume_text: str,
    domain: str,
    embedder: Optional[SentenceTransformer],
) -> float:
    normalized_resume = normalize_text(resume_text)
    domain_skills = extract_skills(resume_text, DOMAIN_SKILL_TERMS[domain])
    target_component = min(1.0, len(domain_skills) / 6)
    core_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Core_NonTech"]))
    software_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Software"]))
    analyst_component = len(extract_skills(resume_text, DOMAIN_SKILL_TERMS["Analyst"]))
    intent_hits = sum(1 for pattern in DOMAIN_INTENT_PATTERNS[domain] if re.search(pattern, normalized_resume))
    intent_component = min(1.0, intent_hits / 4)

    if embedder is not None:
        resume_vec = embedder.encode(normalized_resume, normalize_embeddings=True).reshape(1, -1)
        domain_vec = embedder.encode(DOMAIN_PROTOTYPES[domain], normalize_embeddings=True).reshape(1, -1)
        semantic_component = float(cosine_similarity(resume_vec, domain_vec)[0][0])
    else:
        matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2)).fit_transform(
            [normalized_resume, DOMAIN_PROTOTYPES[domain]]
        )
        semantic_component = float(cosine_similarity(matrix[0], matrix[1])[0][0])

    fit = (0.45 * target_component) + (0.25 * semantic_component) + (0.30 * intent_component)

    if domain == "Core_NonTech" and software_component >= 6 and analyst_component <= 2 and core_component <= 2:
        fit *= 0.20
    if domain == "Software" and software_component >= analyst_component + core_component:
        fit += 0.08
    if domain == "Software" and analyst_component >= 5 and analyst_component >= core_component + 3:
        fit *= 0.88
    if domain == "Analyst" and analyst_component >= core_component:
        fit += 0.04
    if domain == "Analyst" and analyst_component >= 5 and intent_hits >= 2:
        fit += 0.10

    return max(0.0, min(1.0, fit))


def jd_matches_domain(domain: str, jd: Dict) -> bool:
    stored_domain = jd.get("domain")
    stored_conf = jd.get("domain_confidence") or 0.0
    if stored_domain == domain and stored_conf >= 0.55:
        return True
    if jd.get("domain_secondary") == domain and stored_conf >= 0.35:
        return True

    combined = f"{jd['file_name']}\n{jd['cleaned_text'] or jd['jd_text']}"
    counts = {name: len(extract_skills(combined, DOMAIN_SKILL_TERMS[name])) for name in DOMAIN_SKILL_TERMS}
    best_domain = max(counts, key=counts.get)

    if domain == "Software":
        return counts["Software"] >= 2 or (best_domain == "Software" and counts["Software"] >= 1)
    if domain == "Analyst":
        return counts["Analyst"] >= 2 or (best_domain == "Analyst" and counts["Analyst"] >= 1)
    return counts["Core_NonTech"] >= 2 or (
        best_domain == "Core_NonTech" and counts["Core_NonTech"] >= max(counts["Software"], counts["Analyst"])
    )


def compute_ats_score(semantic: float, skill_overlap: float, keyword_density: float) -> float:
    raw = (
        WEIGHT_SEMANTIC * semantic
        + WEIGHT_SKILL * skill_overlap
        + WEIGHT_KEYWORD * keyword_density
    )
    return max(0.0, min(1.0, raw))


def build_reason(overlap: Set[str], semantic: float, skill_score: float, keyword_score: float) -> str:
    parts = []
    if overlap:
        parts.append(f"matched skills: {', '.join(sorted(overlap)[:6])}")
    parts.append("strong semantic match" if semantic >= 0.65 else "moderate semantic match" if semantic >= 0.45 else "low semantic similarity")
    if skill_score < 0.20:
        parts.append("few domain skills in resume")
    if keyword_score < 0.20:
        parts.append("low JD keyword coverage")
    return "; ".join(parts)


def match_resume(
    resume_path: str,
    domain: str,
    top_n: int = 10,
    jd_folder: str = "./extracted/JDs",
    verbose: bool = True,
) -> List[Dict]:
    if domain not in DOMAINS:
        print(f"Unknown domain '{domain}'. Choose from: {list(DOMAINS.keys())}")
        return []

    resume_text = get_text_from_file(resume_path)
    if not resume_text:
        if verbose:
            print("ERROR: Could not extract text from resume.")
        return []

    if verbose:
        print(f"\n{'═' * 60}")
        print(f"  Resume : {resume_path}")
        print(f"  Domain : {domain} — {DOMAINS[domain]}")
        print(f"{'═' * 60}")
        print(f"  Resume text: {len(resume_text):,} chars extracted")

    jds = fetch_jds_by_domain(domain)
    source = "database"
    if not jds:
        if verbose:
            print(f"  No JDs found in DB for '{domain}', falling back to local folder")
        jds = fetch_jds_from_folder(jd_folder)
        source = "local folder"

    if not jds:
        if verbose:
            print("  No JDs found anywhere. Exiting.")
        return []

    embedder = get_model()
    skill_terms = DOMAIN_SKILL_TERMS[domain]
    resume_skills = extract_skills(resume_text, skill_terms)

    if verbose:
        print(f"  JDs fetched: {len(jds)} from {source}")
        print(f"  Encoder  : {'SentenceTransformer (all-MiniLM-L6-v2)' if embedder else 'TF-IDF fallback'}")
        print(f"  Resume skills found: {sorted(resume_skills) or 'none detected'}")
        print("  Scoring …")

    semantic_scores = compute_semantic_scores(resume_text, jds, embedder)
    results = []
    filtered_out = []

    for jd, semantic in zip(jds, semantic_scores):
        if not jd_matches_domain(domain, jd):
            filtered_out.append((jd["file_name"], "domain mismatch"))
            continue
        if semantic < SEMANTIC_RELEVANCE_GATE:
            filtered_out.append((jd["file_name"], f"{round(semantic * 100, 1)}%"))
            continue

        jd_text = jd["jd_text"]
        cleaned = jd["cleaned_text"] or jd_text
        jd_skills = extract_skills(f"{jd['file_name']}\n{cleaned}\n{jd_text}", skill_terms)
        skill_overlap = score_skill_overlap(resume_skills, jd_skills)
        keyword_density = score_keyword_density(resume_text, jd_text)
        ats = compute_ats_score(semantic, skill_overlap, keyword_density)
        overlap_terms = resume_skills & jd_skills

        results.append(
            {
                "file_name": jd["file_name"],
                "ats_score": round(ats * 100, 1),
                "semantic": round(semantic * 100, 1),
                "skill_overlap": round(skill_overlap * 100, 1),
                "keyword_density": round(keyword_density * 100, 1),
                "matched_skills": sorted(overlap_terms),
                "reason": build_reason(overlap_terms, semantic, skill_overlap, keyword_density),
            }
        )

    results.sort(key=lambda item: item["ats_score"], reverse=True)

    if verbose and filtered_out:
        print("\n  Excluded: " + ", ".join(f"{name} ({reason})" for name, reason in filtered_out))

    if not results:
        if verbose:
            print(f"\n  No relevant JDs found for domain [{domain}].")
        return []

    if verbose:
        print(f"\n  Top {min(top_n, len(results))} matches in [{domain}]:\n")
        print(f"  {'#':<3} {'JD File':<42} {'ATS':>5}  {'Sem':>5}  {'Skill':>5}  {'KW':>5}")
        print(f"  {'─' * 70}")
        for rank, item in enumerate(results[:top_n], 1):
            print(
                f"  {rank:<3} {item['file_name'][:42]:<42} "
                f"{item['ats_score']:>5.1f}  {item['semantic']:>5.1f}  "
                f"{item['skill_overlap']:>5.1f}  {item['keyword_density']:>5.1f}"
            )
            print(f"      → {item['reason']}")

        scores = [item["ats_score"] for item in results]
        top = results[0]
        print(f"\n  Scores: min {min(scores):.1f}  max {max(scores):.1f}  avg {sum(scores)/len(scores):.1f}")
        print(f"\n  Best JD: {top['file_name']} ({top['ats_score']}/100)")
        if top["matched_skills"]:
            print(f"  Skills hit: {', '.join(top['matched_skills'][:10])}")
        print(f"\n{'═' * 60}\n")

    return results


def match_resume_all_domains(resume_path: str, top_n_per_domain: int = 3) -> List[Dict]:
    resume_text = get_text_from_file(resume_path)
    embedder = get_model()

    print(f"\n{'═' * 72}")
    print(f"  Resume : {resume_path}")
    print("  Domain-wise ATS summary")
    print(f"{'═' * 72}")

    summary_rows = []
    for domain, label in DOMAINS.items():
        results = match_resume(resume_path, domain, top_n=top_n_per_domain, verbose=False)
        domain_fit = round(compute_resume_domain_fit(resume_text, domain, embedder) * 100, 1)

        if results:
            top_score = results[0]["ats_score"]
            top_file = results[0]["file_name"]
            top_matches = results[: min(3, len(results))]
            avg_top = round(sum(item["ats_score"] for item in top_matches) / len(top_matches), 1)
            hits = len(results)
        else:
            top_score = 0.0
            avg_top = 0.0
            top_file = "-"
            hits = 0

        overall_score = round((DOMAIN_FIT_WEIGHT * domain_fit) + (DOMAIN_BUCKET_WEIGHT * avg_top), 1)
        summary_rows.append(
            {
                "domain": domain,
                "label": label,
                "domain_fit": domain_fit,
                "top_score": top_score,
                "avg_top": avg_top,
                "overall_score": overall_score,
                "hits": hits,
                "top_file": top_file,
            }
        )

    summary_rows.sort(key=lambda item: (item["overall_score"], item["domain_fit"], item["avg_top"]), reverse=True)

    print(f"\n  {'#':<3} {'Domain':<12} {'Fit':>6}  {'Top ATS':>7}  {'Avg Top3':>8}  {'Overall':>7}  {'Hits':>4}  Best JD")
    print(f"  {'─' * 104}")
    for idx, row in enumerate(summary_rows, 1):
        print(
            f"  {idx:<3} {row['domain']:<12} {row['domain_fit']:>6.1f}  {row['top_score']:>7.1f}  "
            f"{row['avg_top']:>8.1f}  {row['overall_score']:>7.1f}  {row['hits']:>4}  "
            f"{row['top_file'][:40]}"
        )

    best = summary_rows[0]
    print(f"\n  Best overall domain: {best['domain']} — {best['label']}")
    print(f"  Resume fit: {best['domain_fit']}/100")
    print(f"  Domain bucket avg top-3 ATS: {best['avg_top']}/100")
    print(f"  Combined overall score: {best['overall_score']}/100")
    print(f"  Best JD in that domain: {best['top_file']} ({best['top_score']}/100)")
    print(f"\n{'═' * 72}\n")
    return summary_rows


if __name__ == "__main__":
    if len(sys.argv) == 3:
        resume_path = sys.argv[1]
        domain = sys.argv[2]
    elif len(sys.argv) == 2:
        resume_path = sys.argv[1]
        domain = None
    else:
        resume_path = input("Resume path (PDF or DOCX): ").strip()
        domain = None

    if not os.path.isfile(resume_path):
        print(f"File not found: {resume_path}")
        sys.exit(1)

    if domain:
        match_resume(resume_path, domain, top_n=10)
    else:
        match_resume_all_domains(resume_path)
